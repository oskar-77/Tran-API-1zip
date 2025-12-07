"""DOCX converter for unified documents."""

import io
import base64
from typing import Optional
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.converters.base_converter import BaseConverter
from src.schemas.document import (
    Document, Page, TextDirection,
    HeadingBlock, ParagraphBlock, ImageBlock, TableBlock, ListBlock, LinkBlock
)


class DocxConverter(BaseConverter):
    """Converter for DOCX output."""
    
    OUTPUT_FORMAT = "docx"
    OUTPUT_EXTENSION = ".docx"
    
    def __init__(self, document: Document, embed_images: bool = True):
        """Initialize DOCX converter.
        
        Args:
            document: Unified document model
            embed_images: Embed images in document
        """
        super().__init__(document)
        self.embed_images = embed_images
        self._doc = None
    
    def convert(self) -> bytes:
        """Convert document to DOCX bytes."""
        self._doc = DocxDocument()
        
        self._setup_styles()
        
        if self.document.title:
            title_para = self._doc.add_heading(self.document.title, level=0)
            self._set_paragraph_direction(title_para, self.document.direction)
        
        for page in self.document.pages:
            self._convert_page(page)
        
        output = io.BytesIO()
        self._doc.save(output)
        output.seek(0)
        
        return output.getvalue()
    
    def save(self, output_path) -> str:
        """Save document to file."""
        from pathlib import Path
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        content = self.convert()
        with open(output_path, 'wb') as f:
            f.write(content)
        
        return str(output_path)
    
    def _setup_styles(self):
        """Setup document styles."""
        styles = self._doc.styles
        
        try:
            normal_style = styles['Normal']
            normal_style.font.name = 'Arial'
            normal_style.font.size = Pt(11)
        except Exception:
            pass
    
    def _convert_page(self, page: Page):
        """Convert a page to DOCX content."""
        if len(self.document.pages) > 1 and page.page_number > 1:
            self._doc.add_page_break()
        
        for block in page.blocks:
            self._convert_block(block, page.direction)
    
    def _convert_block(self, block, page_direction: TextDirection):
        """Convert a block to DOCX content."""
        if isinstance(block, HeadingBlock):
            self._convert_heading(block)
        elif isinstance(block, ParagraphBlock):
            self._convert_paragraph(block)
        elif isinstance(block, ImageBlock):
            self._convert_image(block)
        elif isinstance(block, TableBlock):
            self._convert_table(block)
        elif isinstance(block, ListBlock):
            self._convert_list(block)
        elif isinstance(block, LinkBlock):
            self._convert_link(block)
    
    def _convert_heading(self, block: HeadingBlock):
        """Convert heading block to DOCX."""
        para = self._doc.add_heading(block.text, level=block.level)
        self._set_paragraph_direction(para, block.direction)
        
        if block.links:
            self._add_hyperlinks_to_paragraph(para, block.links)
    
    def _convert_paragraph(self, block: ParagraphBlock):
        """Convert paragraph block to DOCX."""
        para = self._doc.add_paragraph()
        run = para.add_run(block.text)
        
        if block.is_bold:
            run.bold = True
        if block.is_italic:
            run.italic = True
        
        self._set_paragraph_direction(para, block.direction)
        
        if block.links:
            self._add_hyperlinks_to_paragraph(para, block.links)
    
    def _convert_image(self, block: ImageBlock):
        """Convert image block to DOCX."""
        if not self.embed_images:
            if block.caption:
                self._doc.add_paragraph(f"[Image: {block.caption}]")
            return
        
        image_data = None
        
        if block.image_data:
            if block.image_data.startswith('data:'):
                base64_str = block.image_data.split(',')[1]
                image_data = base64.b64decode(base64_str)
            else:
                try:
                    image_data = base64.b64decode(block.image_data)
                except Exception:
                    pass
        
        if image_data:
            try:
                image_stream = io.BytesIO(image_data)
                self._doc.add_picture(image_stream, width=Inches(5))
            except Exception:
                if block.caption:
                    self._doc.add_paragraph(f"[Image: {block.caption}]")
        
        if block.caption:
            caption_para = self._doc.add_paragraph(block.caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.runs[0].italic = True
    
    def _convert_table(self, block: TableBlock):
        """Convert table block to DOCX."""
        if not block.rows:
            return
        
        num_rows = len(block.rows)
        num_cols = max(len(row.cells) for row in block.rows) if block.rows else 0
        
        if num_cols == 0:
            return
        
        table = self._doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row_idx, row in enumerate(block.rows):
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(table.rows[row_idx].cells):
                    table_cell = table.rows[row_idx].cells[col_idx]
                    table_cell.text = cell.content
                    
                    if cell.is_header:
                        for para in table_cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
        
        self._doc.add_paragraph()
    
    def _convert_list(self, block: ListBlock):
        """Convert list block to DOCX."""
        for i, item in enumerate(block.items):
            if block.is_ordered:
                para = self._doc.add_paragraph(item, style='List Number')
            else:
                para = self._doc.add_paragraph(item, style='List Bullet')
            
            self._set_paragraph_direction(para, block.direction)
    
    def _convert_link(self, block: LinkBlock):
        """Convert link block to DOCX."""
        para = self._doc.add_paragraph()
        self._add_hyperlink(para, block.url, block.text)
    
    def _set_paragraph_direction(self, para, direction: TextDirection):
        """Set paragraph text direction."""
        if direction == TextDirection.RTL:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            pPr = para._element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)
    
    def _add_hyperlink(self, paragraph, url: str, text: str):
        """Add a hyperlink to a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        
        hyperlink.append(new_run)
        paragraph._element.append(hyperlink)
    
    def _add_hyperlinks_to_paragraph(self, para, links: list):
        """Add hyperlinks to existing paragraph (placeholder)."""
        pass
